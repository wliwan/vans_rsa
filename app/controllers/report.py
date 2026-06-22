import asyncio
import json
import logging
import os
import re
import time
import uuid
from io import BytesIO
from typing import Dict, List, Optional

import pandas as pd
from openai import OpenAI
from PIL import Image as PILImage

from app.models.admin import AIProxy, AnalysisSheet, Document, OriginalSheet, Report, Skill, StaticFile, Workspace
from app.settings import settings
from app.utils.doc_to_md import file_to_markdown

logger = logging.getLogger(__name__)

# ── 进度状态文件目录 ──
CACHE_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "cache"
)
REPORT_PROGRESS_DIR = os.path.join(CACHE_DIR, "report_generate")
os.makedirs(REPORT_PROGRESS_DIR, exist_ok=True)

# 内存中正在运行的任务缓存（key: task_id）
_running_tasks: Dict[str, dict] = {}


def _progress_file(task_id: str) -> str:
    return os.path.join(REPORT_PROGRESS_DIR, f"{task_id}.json")


def _save_progress(task_id: str, progress: dict):
    with open(_progress_file(task_id), "w", encoding="utf-8") as f:
        json.dump(progress, f, ensure_ascii=False, indent=2)


def _load_progress(task_id: str) -> dict:
    """读取进度：内存优先，文件兜底"""
    if task_id in _running_tasks:
        return dict(_running_tasks[task_id])
    pf = _progress_file(task_id)
    if os.path.exists(pf):
        try:
            with open(pf, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            pass
    return {"status": "not_found", "progress": 0, "phase": "", "message": ""}


def _update_progress(task_id: str, status: str, phase: str, progress: int, message: str, detail: str = ""):
    """更新进度到内存和磁盘"""
    p = {
        "task_id": task_id,
        "status": status,       # running | done | error
        "phase": phase,         # preparing | planning | generating | assembling | saving
        "progress": progress,   # 0-100
        "message": message,
        "detail": detail,
        "started_at": _running_tasks.get(task_id, {}).get("started_at", time.time()),
        "done_at": time.time() if status in ("done", "error") else None,
    }
    _running_tasks[task_id] = p
    _save_progress(task_id, p)


class ReportService:
    # 默认系统提示词（前端弹窗预填，后端作为兜底）
    DEFAULT_SYSTEM_PROMPT = (
        "你是一个专业的数据分析文书生成专家。请根据提供的数据生成一份完整的HTML格式分析文书。\n"
        "文书需包含以下部分（使用HTML标签）：\n"
        "1. <h1> 文书标题\n"
        "2. <h2> 数据概览 — 数据来源、数据量等\n"
        "3. <h2> 核心发现 — 3-5个关键洞察\n"
        "4. <h2> 详细分析 — 多维度数据解读，包含 <table> 表格\n"
        "5. <h2> 结论与建议\n"
        "样式要求：使用内联CSS，简洁专业风格，配色为蓝白灰。"
    )

    # ── AI 调用辅助 ──

    @staticmethod
    def _build_client(ai_proxy) -> OpenAI:
        return OpenAI(base_url=ai_proxy.url, api_key=ai_proxy.token)

    @staticmethod
    async def _call_ai(client: OpenAI, model: str, system_prompt: str, user_prompt: str) -> str:
        def _sync_call():
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.7,
            )
            return response.choices[0].message.content
        return await asyncio.to_thread(_sync_call)

    @staticmethod
    def _format_size(size: int) -> str:
        if size < 1024:
            return f"{size} B"
        elif size < 1024 * 1024:
            return f"{size / 1024:.1f} KB"
        return f"{size / (1024 * 1024):.1f} MB"

    @classmethod
    def _short_link(cls, token: str) -> str:
        base = settings.PUBLIC_BASE_URL.rstrip("/") if settings.PUBLIC_BASE_URL else ""
        return f"{base}/api/sf/{token}"

    @classmethod
    def _generate_image_md(cls, rec) -> str:
        """为图片/静态文件生成元数据 Markdown 描述（不读取二进制内容）"""
        lines = [f"## {rec.name or getattr(rec, 'file_name', '文件')}"]
        if hasattr(rec, 'file_size') and rec.file_size:
            lines.append(f"- 大小: {cls._format_size(rec.file_size)}")
        if hasattr(rec, 'width') and rec.width and hasattr(rec, 'height') and rec.height:
            lines.append(f"- 尺寸: {rec.width}×{rec.height}px")
        if hasattr(rec, 'format_type') and rec.format_type:
            lines.append(f"- 格式: {rec.format_type}")
        if hasattr(rec, 'color_mode') and rec.color_mode:
            lines.append(f"- 色彩: {rec.color_mode}")
        if hasattr(rec, 'dpi') and rec.dpi:
            lines.append(f"- DPI: {rec.dpi}")
        if hasattr(rec, 'bit_depth') and rec.bit_depth:
            lines.append(f"- 位深: {rec.bit_depth}bit")
        if hasattr(rec, 'short_url_token') and rec.short_url_token:
            lines.append(f"- 短链接: `{cls._short_link(rec.short_url_token)}`")
        if hasattr(rec, 'description') and rec.description:
            lines.append(f"- 描述: {rec.description}")
        if hasattr(rec, 'source') and rec.source:
            source_labels = {
                "upload": "用户上传", "cv_processed": "CV处理",
                "ai_generated": "AI生成", "ocr": "OCR提取",
                "road_material": "路网素材", "auto": "自动生成",
            }
            lines.append(f"- 来源: {source_labels.get(rec.source, rec.source)}")
        return "\n".join(lines)

    @classmethod
    async def _read_all_data(cls, sheet_ids: List[int], analysis_ids: List[int], document_ids: List[int] = None, static_file_ids: List[int] = None) -> str:
        """异步读取所有数据源：表格（Excel/CSV） + 文档（PDF/DOCX/MD等） + 静态文件（图片元数据）"""
        if document_ids is None:
            document_ids = []
        if static_file_ids is None:
            static_file_ids = []

        tasks = []
        for sid in sheet_ids:
            tasks.append(OriginalSheet.get_or_none(id=sid))
        for aid in analysis_ids:
            tasks.append(AnalysisSheet.get_or_none(id=aid))
        for did in document_ids:
            tasks.append(Document.get_or_none(id=did))
        for sfid in static_file_ids:
            tasks.append(StaticFile.get_or_none(id=sfid))
        records = await asyncio.gather(*tasks)

        def _read_one(rec):
            if rec is None:
                return None
            try:
                rec_name = getattr(rec, 'name', None) or getattr(rec, 'file_name', '文件')

                is_static_image = hasattr(rec, 'is_image') and rec.is_image
                if is_static_image:
                    ext = os.path.splitext(getattr(rec, 'file_name', '') or getattr(rec, 'file_path', ''))[1].lower()
                    if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif', '.svg'):
                        return ReportService._generate_image_md(rec)

                file_path = getattr(rec, 'file_path', None)
                if not file_path:
                    return None

                if not os.path.exists(file_path):
                    logger.warning(f"文件不存在: {file_path}")
                    return f"## {rec_name}\n\n*(文件不存在: {os.path.basename(file_path)})*"

                ext = os.path.splitext(file_path)[1].lower()
                ext_name = os.path.basename(file_path)

                if ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif'):
                    return ReportService._generate_image_md(rec)

                if ext == '.csv':
                    df = pd.read_csv(file_path)
                    return f"## {rec_name}\n\n{df.head(80).to_markdown(index=False)}"

                if ext in ('.xlsx', '.xls'):
                    df = pd.read_excel(file_path)
                    return f"## {rec_name}\n\n{df.head(80).to_markdown(index=False)}"

                try:
                    md_text, _ = file_to_markdown(file_path, ext_name)
                    return f"## {rec_name}\n\n{md_text[:5000]}"
                except ValueError:
                    try:
                        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                            text = f.read()
                        return f"## {rec_name}\n\n{text[:5000]}"
                    except Exception:
                        logger.warning(f"无法读取文件内容（二进制/未知格式）: {file_path}")
                        return ReportService._generate_image_md(rec)
            except Exception as e:
                logger.warning(f"读取文件失败 {getattr(rec, 'file_path', '?')}: {e}")
                return None

        parts = []
        for rec in records:
            result = await asyncio.to_thread(_read_one, rec)
            if result:
                parts.append(result)
        return "\n\n".join(parts) if parts else "（无数据源）"

    # ═══════════════════════════════════════════════════════
    # 异步多阶段生成（新版）
    # ═══════════════════════════════════════════════════════

    @classmethod
    async def start_generate_report(
        cls,
        workspace_id: int,
        name: str,
        sheet_ids: List[int],
        analysis_ids: List[int],
        ai_proxy_id: int,
        skill_id: Optional[int] = None,
        extra_prompt: str = "",
        system_prompt: str = "",
        document_ids: List[int] = None,
        static_file_ids: List[int] = None,
    ) -> str:
        """
        启动异步文书生成，返回 task_id。
        前端通过轮询 GET /report/generate-progress?task_id=xxx 获取进度，
        完成后通过 GET /report/generate-result?task_id=xxx 获取生成的文书。
        """
        if document_ids is None:
            document_ids = []
        if static_file_ids is None:
            static_file_ids = []

        task_id = uuid.uuid4().hex[:12]
        _update_progress(task_id, "running", "preparing", 0, "任务已创建，正在准备...")

        # 后台启动异步任务
        asyncio.create_task(
            cls._do_generate_report(
                task_id=task_id,
                workspace_id=workspace_id,
                name=name,
                sheet_ids=sheet_ids,
                analysis_ids=analysis_ids,
                ai_proxy_id=ai_proxy_id,
                skill_id=skill_id,
                extra_prompt=extra_prompt,
                system_prompt=system_prompt,
                document_ids=document_ids,
                static_file_ids=static_file_ids,
            )
        )

        return task_id

    @classmethod
    async def _do_generate_report(
        cls,
        task_id: str,
        workspace_id: int,
        name: str,
        sheet_ids: List[int],
        analysis_ids: List[int],
        ai_proxy_id: int,
        skill_id: Optional[int] = None,
        extra_prompt: str = "",
        system_prompt: str = "",
        document_ids: List[int] = None,
        static_file_ids: List[int] = None,
    ):
        """后台单AI文书生成（由 asyncio.create_task 调度，保留异步+进度反馈）"""
        try:
            # ── 阶段 1: 读取素材 & Skill (0-20%) ──
            _update_progress(task_id, "running", "preparing", 2, "正在读取素材数据...")
            total_ids = len(sheet_ids) + len(analysis_ids) + len(document_ids) + len(static_file_ids)

            skill_prompt = ""
            if skill_id:
                skill = await Skill.filter(id=skill_id).first()
                if skill:
                    skill_prompt = skill.content
                    _update_progress(task_id, "running", "preparing", 5, f"已加载 Skill: {skill.title}")

            effective_system_prompt = system_prompt.strip() if system_prompt and system_prompt.strip() else cls.DEFAULT_SYSTEM_PROMPT

            _update_progress(task_id, "running", "preparing", 10, f"正在读取 {total_ids} 个数据源...")
            data_text = await cls._read_all_data(sheet_ids, analysis_ids, document_ids, static_file_ids)
            _update_progress(task_id, "running", "preparing", 20, f"素材读取完成（共 {total_ids} 项）")

            # ── 阶段 2: 单 AI 调用生成文书 (20-90%) ──
            _update_progress(task_id, "running", "generating", 25, "AI 正在生成文书...")

            ai_proxy = await AIProxy.get(id=ai_proxy_id)
            client = cls._build_client(ai_proxy)
            model = ai_proxy.model or "gpt-3.5-turbo"

            if skill_prompt:
                effective_system_prompt += f"\n\n文书框架参考：\n{skill_prompt}"

            user_prompt = f"文书标题：{name}\n\n数据来源：\n{data_text}"
            if extra_prompt:
                user_prompt += f"\n\n额外要求：{extra_prompt}"
            user_prompt += "\n\n请生成完整的HTML文书（包含<!DOCTYPE html>声明），直接返回HTML代码。"

            # 发起 AI 调用，同时用模拟心跳更新进度（因为单次 AI 调用期间无法获取真实进度）
            _update_progress(task_id, "running", "generating", 30, "AI 正在分析素材并撰写文书...")

            html_content = await cls._call_ai(client, model, effective_system_prompt, user_prompt)
            _update_progress(task_id, "running", "generating", 85, "AI 生成完成，正在清洗格式...")

            html_content = cls._clean_html(html_content)
            _update_progress(task_id, "running", "generating", 90, "文书格式清洗完成")

            # ── 阶段 3: 保存到数据库 (90-100%) ──
            _update_progress(task_id, "running", "saving", 93, "正在保存文书...")

            report = await Report.create(
                workspace_id=workspace_id,
                name=name,
                content=html_content,
                source_sheet_ids=sheet_ids,
                source_analysis_ids=analysis_ids,
                source_document_ids=document_ids,
                source_static_ids=static_file_ids,
                ai_proxy_id=ai_proxy_id,
                skill_id=skill_id,
                prompt=extra_prompt,
                system_prompt=system_prompt or cls.DEFAULT_SYSTEM_PROMPT,
            )

            result_data = await report.to_dict()
            _update_progress(task_id, "done", "saving", 100, "文书生成完成",
                             detail=json.dumps({"report_id": report.id, "name": name}, ensure_ascii=False))
            # 保存结果到进度文件
            p = _load_progress(task_id)
            p["result"] = result_data
            _save_progress(task_id, p)

            logger.info(f"文书生成完成: task_id={task_id}, report_id={report.id}, name={name}")

        except Exception as e:
            logger.exception(f"文书生成失败: task_id={task_id}")
            _update_progress(task_id, "error", "error", 0, "文书生成失败", detail=str(e))

    @classmethod
    def get_generate_progress(cls, task_id: str) -> dict:
        """获取文书生成进度"""
        return _load_progress(task_id)

    @classmethod
    async def get_generate_result(cls, task_id: str) -> Optional[dict]:
        """获取已生成的文书结果，完成后自动清理进度文件"""
        progress = _load_progress(task_id)
        if progress.get("status") != "done":
            return None
        result = progress.get("result")
        if not result:
            # 尝试从数据库读取
            detail_str = progress.get("detail", "")
            try:
                detail = json.loads(detail_str) if detail_str else {}
                report_id = detail.get("report_id")
                if report_id:
                    report = await Report.get_or_none(id=report_id)
                    if report:
                        result = await report.to_dict()
            except (json.JSONDecodeError, Exception):
                pass

        # 清理进度文件
        try:
            pf = _progress_file(task_id)
            if os.path.exists(pf):
                os.remove(pf)
        except Exception:
            pass
        _running_tasks.pop(task_id, None)

        return result

    @staticmethod
    def _clean_html(text: str) -> str:
        """从 AI 返回内容中提取纯 HTML 文档，丢弃所有非 HTML 的说明文本。"""
        text = text.strip()
        if text.startswith("```html"):
            text = text[7:]
        elif text.startswith("```"):
            text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()

        html_start = -1
        html_end = -1

        doctype_idx = text.lower().find("<!doctype")
        html_tag_idx = text.lower().find("<html")
        if doctype_idx != -1:
            html_start = doctype_idx
        elif html_tag_idx != -1:
            html_start = html_tag_idx

        if html_start != -1:
            end_idx = text.lower().rfind("</html>")
            if end_idx != -1:
                html_end = end_idx + len("</html>")

        if html_start != -1 and html_end != -1:
            text = text[html_start:html_end]

        if text.strip().lower().startswith("<html"):
            text = "<!DOCTYPE html>\n" + text

        return text.strip()

    # ═══════════════════════════════════════════════════════
    # 以下为旧版同步生成方法（保留兼容）
    # ═══════════════════════════════════════════════════════

    @classmethod
    async def generate_report(
        cls,
        workspace_id: int,
        name: str,
        sheet_ids: List[int],
        analysis_ids: List[int],
        ai_proxy_id: int,
        skill_id: Optional[int] = None,
        extra_prompt: str = "",
        system_prompt: str = "",
        document_ids: List[int] = None,
        static_file_ids: List[int] = None,
    ) -> Report:
        """旧版同步生成方法（保留向后兼容）"""
        if document_ids is None:
            document_ids = []
        if static_file_ids is None:
            static_file_ids = []
        ai_proxy = await AIProxy.get(id=ai_proxy_id)
        skill_prompt = ""
        if skill_id:
            skill = await Skill.filter(id=skill_id).first()
            if skill:
                skill_prompt = skill.content

        data_text = await cls._read_all_data(sheet_ids, analysis_ids, document_ids, static_file_ids)

        effective_system_prompt = system_prompt.strip() if system_prompt and system_prompt.strip() else cls.DEFAULT_SYSTEM_PROMPT

        if skill_prompt:
            effective_system_prompt += f"\n\n文书框架参考：\n{skill_prompt}"

        user_prompt = f"文书标题：{name}\n\n数据来源：\n{data_text}"
        if extra_prompt:
            user_prompt += f"\n\n额外要求：{extra_prompt}"
        user_prompt += "\n\n请生成完整的HTML文书（包含<!DOCTYPE html>声明），直接返回HTML代码。"

        client = cls._build_client(ai_proxy)
        html_content = await cls._call_ai(client, ai_proxy.model or "gpt-3.5-turbo", effective_system_prompt, user_prompt)
        html_content = cls._clean_html(html_content)

        report = await Report.create(
            workspace_id=workspace_id,
            name=name,
            content=html_content,
            source_sheet_ids=sheet_ids,
            source_analysis_ids=analysis_ids,
            source_document_ids=document_ids,
            source_static_ids=static_file_ids,
            ai_proxy_id=ai_proxy_id,
            skill_id=skill_id,
            prompt=extra_prompt,
            system_prompt=system_prompt or cls.DEFAULT_SYSTEM_PROMPT,
        )
        return report

    # ═══════════════════════════════════════════════════════
    # 导出方法
    # ═══════════════════════════════════════════════════════

    @classmethod
    async def export_html(cls, report_id: int) -> str:
        report = await Report.get(id=report_id)
        return report.content

    @classmethod
    async def export_pdf_bytes(cls, report_id: int) -> bytes:
        try:
            from weasyprint import HTML
            report = await Report.get(id=report_id)
            html = HTML(string=report.content)
            return html.write_pdf()
        except ImportError:
            raise RuntimeError("weasyprint 未安装")

    @classmethod
    async def export_docx_bytes(cls, report_id: int) -> bytes:
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            report = await Report.get(id=report_id)
            soup = BeautifulSoup(report.content, "html.parser")

            doc = Document()
            for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "table", "ul", "ol"]):
                if tag.name == "h1":
                    doc.add_heading(tag.get_text(), level=1)
                elif tag.name == "h2":
                    doc.add_heading(tag.get_text(), level=2)
                elif tag.name == "h3":
                    doc.add_heading(tag.get_text(), level=3)
                elif tag.name in ("p",):
                    doc.add_paragraph(tag.get_text())
                elif tag.name == "table":
                    rows = tag.find_all("tr")
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["td", "th"])))
                        for i, row in enumerate(rows):
                            for j, cell in enumerate(row.find_all(["td", "th"])):
                                table.cell(i, j).text = cell.get_text()
                elif tag.name in ("ul", "ol"):
                    for li in tag.find_all("li"):
                        doc.add_paragraph(li.get_text(), style="List Bullet")

            buf = BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except ImportError:
            raise RuntimeError("python-docx 未安装")

    # ═══════════════════════════════════════════════════════
    # 数据源预览
    # ═══════════════════════════════════════════════════════

    @staticmethod
    def _estimate_tokens(text: str) -> int:
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_words = len(re.findall(r'[a-zA-Z]+', text))
        other = max(0, len(text) - chinese_chars - sum(len(w) for w in re.findall(r'[a-zA-Z]+', text)))
        return int(chinese_chars * 1.5 + english_words * 1.3 + other * 0.3)

    @classmethod
    async def preview_sources(
        cls,
        sheet_ids: List[int],
        analysis_ids: List[int],
        document_ids: List[int],
        static_file_ids: List[int],
    ) -> dict:
        items = []
        md_parts = []

        # 1. 原始表格
        for sid in sheet_ids:
            sheet = await OriginalSheet.get_or_none(id=sid)
            if not sheet:
                continue
            try:
                ext = os.path.splitext(sheet.file_path)[1].lower() if sheet.file_path else ''
                if ext in ('.xlsx', '.xls', '.csv'):
                    df = pd.read_excel(sheet.file_path) if ext != '.csv' else pd.read_csv(sheet.file_path)
                    md_table = df.head(30).to_markdown(index=False)
                    preview = f"**{sheet.name}** ({cls._format_size(sheet.file_size or 0)})\n\n{md_table}\n\n*共 {len(df)} 行 × {len(df.columns)} 列*"
                else:
                    try:
                        text, _ = file_to_markdown(sheet.file_path, os.path.basename(sheet.file_path))
                        preview = f"**{sheet.name}** ({cls._format_size(sheet.file_size or 0)})\n\n{text[:3000]}"
                    except (ValueError, Exception):
                        try:
                            with open(sheet.file_path, 'r', encoding='utf-8', errors='replace') as f:
                                text = f.read()[:3000]
                        except Exception:
                            text = f"*（无法预览）*"
                        preview = f"**{sheet.name}** ({cls._format_size(sheet.file_size or 0)})\n\n```\n{text}\n```"
                items.append({
                    "id": sheet.id, "name": sheet.name, "type": "sheet",
                    "preview": preview, "char_count": len(preview),
                    "file_size": sheet.file_size or 0,
                })
                md_parts.append(preview)
            except Exception as e:
                logger.warning(f"预览表格失败 {sheet.file_path}: {e}")

        # 2. 分析表格
        for aid in analysis_ids:
            analysis = await AnalysisSheet.get_or_none(id=aid)
            if not analysis:
                continue
            try:
                ext = os.path.splitext(analysis.file_path)[1].lower() if analysis.file_path else ''
                if ext in ('.xlsx', '.xls', '.csv'):
                    df = pd.read_excel(analysis.file_path) if ext != '.csv' else pd.read_csv(analysis.file_path)
                    md_table = df.head(30).to_markdown(index=False)
                    preview = f"**{analysis.name}** ({cls._format_size(analysis.file_size or 0)})\n\n{md_table}\n\n*共 {len(df)} 行 × {len(df.columns)} 列*"
                else:
                    try:
                        text, _ = file_to_markdown(analysis.file_path, os.path.basename(analysis.file_path))
                        preview = f"**{analysis.name}** ({cls._format_size(analysis.file_size or 0)})\n\n{text[:3000]}"
                    except (ValueError, Exception):
                        try:
                            with open(analysis.file_path, 'r', encoding='utf-8', errors='replace') as f:
                                text = f.read()[:3000]
                        except Exception:
                            text = f"*（无法预览）*"
                        preview = f"**{analysis.name}** ({cls._format_size(analysis.file_size or 0)})\n\n```\n{text}\n```"
                items.append({
                    "id": analysis.id, "name": analysis.name, "type": "analysis",
                    "preview": preview, "char_count": len(preview),
                    "file_size": analysis.file_size or 0,
                })
                md_parts.append(preview)
            except Exception as e:
                logger.warning(f"预览分析表格失败 {analysis.file_path}: {e}")

        # 3. 文档
        for did in document_ids:
            doc = await Document.get_or_none(id=did)
            if not doc:
                continue
            try:
                if doc.file_path and os.path.exists(doc.file_path):
                    ext = os.path.splitext(doc.file_path)[1].lower()
                    ext_name = os.path.basename(doc.file_path)
                    if ext in ('.xlsx', '.xls', '.csv'):
                        df = pd.read_excel(doc.file_path) if ext != '.csv' else pd.read_csv(doc.file_path)
                        text = df.head(50).to_markdown(index=False)
                    elif ext in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.tiff', '.tif'):
                        text = f"*（图片文件: {ext_name}）*"
                    else:
                        try:
                            text, _ = file_to_markdown(doc.file_path, ext_name)
                            text = text[:5000]
                        except ValueError:
                            try:
                                with open(doc.file_path, 'r', encoding='utf-8', errors='replace') as f:
                                    text = f.read()[:5000]
                            except Exception:
                                text = f"*（无法预览: {ext_name}）*"
                else:
                    text = f"*（空文档）*"
                preview_header = f"**{doc.name}**"
                if doc.file_size:
                    preview_header += f" ({cls._format_size(doc.file_size)})"
                if doc.char_count:
                    preview_header += f" — {doc.char_count} 字"
                if doc.import_source and doc.import_source != 'upload':
                    preview_header += f" — 来源: {doc.import_source}"
                if doc.row_count:
                    preview_header += f" — {doc.row_count} 行"
                preview = f"{preview_header}\n\n{text}"
                items.append({
                    "id": doc.id, "name": doc.name, "type": "document",
                    "preview": preview, "char_count": len(preview),
                    "file_size": doc.file_size or 0,
                })
                md_parts.append(preview)
            except Exception as e:
                logger.warning(f"预览文档失败 {doc.file_path}: {e}")

        # 4. 静态文件
        for sfid in static_file_ids:
            sf = await StaticFile.get_or_none(id=sfid)
            if not sf:
                continue
            try:
                lines = [f"**{sf.name or sf.file_name}**"]
                if sf.file_size:
                    lines.append(f"- 大小: {cls._format_size(sf.file_size)}")
                if sf.short_url_token:
                    lines.append(f"- 链接: `{cls._short_link(sf.short_url_token)}`")
                if sf.is_image and sf.width and sf.height:
                    lines.append(f"- 尺寸: {sf.width}×{sf.height}px")
                if sf.format_type:
                    lines.append(f"- 格式: {sf.format_type}")
                if sf.color_mode:
                    lines.append(f"- 色彩: {sf.color_mode}")
                if sf.dpi:
                    lines.append(f"- DPI: {sf.dpi}")
                if sf.bit_depth:
                    lines.append(f"- 位深: {sf.bit_depth}bit")
                if sf.description:
                    lines.append(f"- 描述: {sf.description}")
                if sf.source:
                    source_labels = {
                        "upload": "用户上传", "cv_processed": "CV处理",
                        "ai_generated": "AI生成", "ocr": "OCR提取",
                        "road_material": "路网素材", "auto": "自动生成",
                    }
                    lines.append(f"- 来源: {source_labels.get(sf.source, sf.source)}")
                if sf.exif_data:
                    exif_preview = {k: v for k, v in sf.exif_data.items() if v}
                    if exif_preview:
                        lines.append(f"- EXIF: {exif_preview}")
                preview = "\n".join(lines)
                items.append({
                    "id": sf.id, "name": sf.name or sf.file_name, "type": "static_file",
                    "preview": preview, "char_count": len(preview),
                    "file_size": sf.file_size or 0,
                })
                md_parts.append(preview)
            except Exception as e:
                logger.warning(f"预览静态文件失败 {sf.file_path}: {e}")

        md_full = "\n\n---\n\n".join(md_parts) if md_parts else "（无数据源）"
        total_chars = sum(it["char_count"] for it in items)
        estimated_tokens = cls._estimate_tokens(md_full)

        return {
            "md_preview": md_full,
            "items": items,
            "total_chars": total_chars,
            "estimated_tokens": estimated_tokens,
            "source_counts": {
                "sheets": len([it for it in items if it["type"] == "sheet"]),
                "analyses": len([it for it in items if it["type"] == "analysis"]),
                "documents": len([it for it in items if it["type"] == "document"]),
                "static_files": len([it for it in items if it["type"] == "static_file"]),
                "total": len(items),
            },
        }


report_service = ReportService()
